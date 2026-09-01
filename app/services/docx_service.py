import re
import xml.sax.saxutils as saxutils
from pathlib import Path
from datetime import datetime
import json

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate


OUTPUT_DIR = Path("applications")
OUTPUT_DIR.mkdir(exist_ok=True)

_BOLD_SPLIT = re.compile(r"(\*\*.+?\*\*)")
_BODY_FONT = "Calibri"
_MUTED_RGB = RGBColor(0x40, 0x40, 0x40)
_META_RGB = RGBColor(0x55, 0x55, 0x55)
_MUTED_HEX = colors.HexColor("#404040")
_META_HEX = colors.HexColor("#555555")
_RULE_HEX = colors.HexColor("#999999")


def _apply_compact_paragraph_spacing(document):
    """python-docx's built-in blank-document template carries Word's default
    ~10pt space-after + ~1.15 line spacing on every paragraph, plus 1in/
    1.25in page margins -- reasonable for essay prose, but it turns a
    dense, mostly single-line-bullet resume into several pages of
    whitespace rather than content (Task 21.13 section 5). Tightening
    spacing/margins to common professional-resume conventions lets the
    rendered page count reflect actual content density instead of template
    padding, without removing or shortening any actual evidence.

    Task 21.31: also sets the base executive typeface/size here (Calibri,
    a font reliably present alongside any DOCX-capable install) and a
    hanging indent on bullet styles so a wrapped bullet line aligns under
    the bullet's TEXT, not its glyph -- the two remaining ingredients for
    an executive-quality look that apply uniformly regardless of which
    named-style paragraphs (_add_*_paragraph below) end up used."""
    for style_name in ("Normal", "List Bullet", "List Bullet 2", "List Bullet 3"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = _BODY_FONT
        style.font.size = Pt(10.5)
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(3)
        paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Bullet 2", "List Bullet 3"):
        style = document.styles[style_name]
        style.paragraph_format.left_indent = Pt(18)
        style.paragraph_format.first_line_indent = Pt(-13)

    for level in (1, 2, 3):
        try:
            style = document.styles[f"Heading {level}"]
        except KeyError:
            continue
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(10 if level == 1 else 6)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = 1.0

    for section in document.sections:
        section.top_margin = Pt(50)     # ~0.7"
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(54)    # 0.75"
        section.right_margin = Pt(54)


def _add_bottom_border(paragraph):
    """A thin, understated rule under a section heading -- the "optional
    simple horizontal separator" the resume design standard allows.  Pure
    paragraph-border XML, no drawing object/text box, so it never affects
    ATS text extraction."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "999999")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_inline_runs(paragraph, text):
    """Split `**bold**` markers out of a line and add real bold/plain runs."""

    for segment in _BOLD_SPLIT.split(text):

        if not segment:
            continue

        if segment.startswith("**") and segment.endswith("**") and len(segment) > 4:
            paragraph.add_run(segment[2:-2]).bold = True
        else:
            paragraph.add_run(segment)


def _classify_markdown_lines(markdown_text):
    """Shared structural pass over the Markdown subset ResumeGenerator/
    CoverLetterGenerator actually produce, so the DOCX and PDF renderers
    stay materially consistent (Task 21.31 section 4) instead of each
    re-implementing their own reading of "what is this line".

    Yields (kind, text) pairs. kind is one of: name (the H1 header line --
    always the candidate's name), headline (a whole-line **bold** directly
    under the name -- the professional headline, when present), contact
    (the plain contact-details line directly under the name/headline),
    section (## heading), entry (### heading -- a job/education/project
    title line), meta (the plain line immediately after an entry heading --
    ResumeGenerator always emits period/location there), bullet ("- "),
    body (anything else). Never invents or reorders content -- purely a
    read of the line order ResumeGenerator/CoverLetterGenerator already
    guarantee.
    """
    in_header_zone = False
    headline_done = False
    contact_done = False
    prev_was_entry = False

    for raw_line in markdown_text.split("\n"):
        stripped = raw_line.strip()

        if not stripped or stripped == "---":
            prev_was_entry = False
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1:
                in_header_zone, headline_done, contact_done = True, False, False
                yield ("name", text)
            elif level == 2:
                in_header_zone = False
                yield ("section", text)
            else:
                yield ("entry", text)
            prev_was_entry = level == 3
            continue

        if stripped.startswith("- "):
            prev_was_entry = False
            yield ("bullet", stripped[2:])
            continue

        if in_header_zone and not headline_done and stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            headline_done = True
            prev_was_entry = False
            yield ("headline", stripped[2:-2])
            continue

        if in_header_zone and not contact_done:
            contact_done = True
            prev_was_entry = False
            yield ("contact", stripped)
            continue

        if prev_was_entry:
            prev_was_entry = False
            yield ("meta", stripped)
            continue

        prev_was_entry = False
        yield ("body", stripped)


def _write_markdown_to_docx(document, markdown_text):
    """Render simple resume/cover-letter Markdown as real, executive-quality
    Word formatting -- single-column, no tables/text boxes/images, every
    heading/date/bullet a real styled paragraph so ATS text extraction sees
    normal readable text throughout. Never writes literal Markdown syntax
    or raw Python dict/list reprs into the document, and never alters the
    underlying candidate facts -- only how they are laid out.
    """
    for kind, text in _classify_markdown_lines(markdown_text):
        if kind == "name":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = _BODY_FONT
            run.font.size = Pt(20)
        elif kind == "headline":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = _BODY_FONT
            run.font.size = Pt(12)
            run.font.color.rgb = _MUTED_RGB
        elif kind == "contact":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            run = paragraph.add_run(text)
            run.font.name = _BODY_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = _MUTED_RGB
        elif kind == "section":
            _add_section_heading_docx(document, text)
        elif kind == "entry":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(9)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.keep_with_next = True
            _add_inline_runs(paragraph, text)
            for run in paragraph.runs:
                run.bold = True
                run.font.name = _BODY_FONT
                run.font.size = Pt(10.5)
        elif kind == "meta":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(text)
            run.italic = True
            run.font.name = _BODY_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = _META_RGB
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, text)
        else:  # body
            paragraph = document.add_paragraph()
            _add_inline_runs(paragraph, text)


def _add_section_heading_docx(document, text):
    """A major section heading (## in Markdown, or the cover letter's own
    "Cover Letter" label): bold, restrained caps, a subtle rule beneath --
    the ALL-CAPS display is a pure character-formatting transform (the
    underlying run text is uppercased directly, matching the PDF sibling
    exactly per Task 21.31 section 4) applied only to structural section
    labels, never to a candidate fact."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = Pt(11.5)
    _add_bottom_border(paragraph)
    return paragraph


def _pdf_styles():
    """A plain, single-column executive style sheet: real embedded/
    selectable text, no tables, images, or text boxes -- kept ATS-parseable
    (Task 21.30 Section 1), and matched to the DOCX renderer's hierarchy/
    sizing one-for-one (Task 21.31 section 4) rather than a separate,
    drifting design."""
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ResumeBody", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10.5, leading=14, spaceAfter=3, alignment=TA_LEFT,
    ))
    styles.add(ParagraphStyle(
        name="ResumeBullet", parent=styles["ResumeBody"], leftIndent=16, bulletIndent=2, spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name="ResumeName", parent=styles["ResumeBody"], fontName="Helvetica-Bold",
        fontSize=20, leading=23, spaceAfter=2, keepWithNext=True,
    ))
    styles.add(ParagraphStyle(
        name="ResumeHeadline", parent=styles["ResumeBody"], fontName="Helvetica-Bold",
        fontSize=12, leading=15, spaceAfter=4, textColor=_MUTED_HEX, keepWithNext=True,
    ))
    styles.add(ParagraphStyle(
        name="ResumeContact", parent=styles["ResumeBody"], fontName="Helvetica",
        fontSize=9.5, leading=12, spaceAfter=12, textColor=_MUTED_HEX,
    ))
    styles.add(ParagraphStyle(
        name="ResumeSection", parent=styles["ResumeBody"], fontName="Helvetica-Bold",
        fontSize=11.5, leading=14, spaceBefore=14, spaceAfter=2, keepWithNext=True,
    ))
    styles.add(ParagraphStyle(
        name="ResumeEntry", parent=styles["ResumeBody"], fontName="Helvetica-Bold",
        fontSize=10.5, leading=13, spaceBefore=9, spaceAfter=1, keepWithNext=True,
    ))
    styles.add(ParagraphStyle(
        name="ResumeMeta", parent=styles["ResumeBody"], fontName="Helvetica-Oblique",
        fontSize=9.5, leading=12, spaceAfter=4, textColor=_META_HEX, keepWithNext=True,
    ))
    return styles


def _section_rule():
    """The PDF sibling of _add_bottom_border: a thin vector line, not an
    image, so it never affects text extraction."""
    return HRFlowable(width="100%", thickness=0.6, color=_RULE_HEX, spaceBefore=1, spaceAfter=6)


def _inline_markup(text):
    """Escape XML-special characters, then re-apply `**bold**` as reportlab's
    own minimal `<b>` markup -- mirrors _add_inline_runs's docx behaviour
    without ever passing unescaped candidate text into the PDF's markup
    parser."""
    parts = []
    for segment in _BOLD_SPLIT.split(text):
        if not segment:
            continue
        if segment.startswith("**") and segment.endswith("**") and len(segment) > 4:
            parts.append(f"<b>{saxutils.escape(segment[2:-2])}</b>")
        else:
            parts.append(saxutils.escape(segment))
    return "".join(parts)


def _write_markdown_to_pdf_story(markdown_text, styles):
    """Same Markdown subset, same _classify_markdown_lines structural read,
    and the same visual hierarchy as _write_markdown_to_docx -- same
    factual content, no images, materially consistent with the DOCX
    sibling (Task 21.31 section 4)."""
    story = []
    for kind, text in _classify_markdown_lines(markdown_text):
        if kind == "name":
            story.append(Paragraph(_inline_markup(text), styles["ResumeName"]))
        elif kind == "headline":
            story.append(Paragraph(_inline_markup(text), styles["ResumeHeadline"]))
        elif kind == "contact":
            story.append(Paragraph(_inline_markup(text), styles["ResumeContact"]))
        elif kind == "section":
            story.append(Paragraph(_inline_markup(text.upper()), styles["ResumeSection"]))
            story.append(_section_rule())
        elif kind == "entry":
            story.append(Paragraph(_inline_markup(text), styles["ResumeEntry"]))
        elif kind == "meta":
            story.append(Paragraph(_inline_markup(text), styles["ResumeMeta"]))
        elif kind == "bullet":
            story.append(Paragraph("&bull;&nbsp;&nbsp;" + _inline_markup(text), styles["ResumeBullet"]))
        else:  # body
            story.append(Paragraph(_inline_markup(text), styles["ResumeBody"]))
    return story


def _build_pdf(filename, story):
    document = SimpleDocTemplate(
        str(filename), pagesize=letter,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )
    document.build(story)


def _create_application_folder(company, job_title):

    company = (company or "").replace("/", "-").strip()

    job_title = (job_title or "Unknown Position").replace("/", "-").strip()

    if not company:
        company = "Unknown Company"

    folder = OUTPUT_DIR / f"{company}_{job_title}"

    folder.mkdir(parents=True, exist_ok=True)

    return folder


def generate_resume_docx(
    resume_text,
    company="Company",
    job_title="Position"
):

    folder = _create_application_folder(
        company,
        job_title
    )

    filename = folder / "Resume.docx"

    document = Document()
    _apply_compact_paragraph_spacing(document)

    _write_markdown_to_docx(document, resume_text)

    document.save(filename)

    return str(filename)


def generate_resume_pdf(
    resume_text,
    company="Company",
    job_title="Position"
):
    """Text-based, ATS-safe PDF sibling of generate_resume_docx, rendered
    from the SAME already-approved resume content (Task 21.30 Section 1) --
    never a re-generation, never an image/rasterized conversion."""

    folder = _create_application_folder(
        company,
        job_title
    )

    filename = folder / "Resume.pdf"

    story = _write_markdown_to_pdf_story(resume_text, _pdf_styles())
    _build_pdf(filename, story)

    return str(filename)


def generate_cover_letter_docx(
    cover_letter,
    company="Company",
    job_title="Position"
):

    folder = _create_application_folder(
        company,
        job_title
    )

    filename = folder / "CoverLetter.docx"

    document = Document()
    _apply_compact_paragraph_spacing(document)

    _add_section_heading_docx(document, "Cover Letter")

    _write_markdown_to_docx(document, cover_letter)

    document.save(filename)

    return str(filename)


def generate_cover_letter_pdf(
    cover_letter,
    company="Company",
    job_title="Position"
):
    """Text-based, ATS-safe PDF sibling of generate_cover_letter_docx, from
    the SAME already-approved cover-letter content (Task 21.30 Section 1)."""

    folder = _create_application_folder(
        company,
        job_title
    )

    filename = folder / "CoverLetter.pdf"

    styles = _pdf_styles()
    story = [Paragraph("COVER LETTER", styles["ResumeSection"]), _section_rule()]
    story.extend(_write_markdown_to_pdf_story(cover_letter, styles))
    _build_pdf(filename, story)

    return str(filename)


def save_career_report(
    decision,
    employer,
    company,
    job_title
):

    folder = _create_application_folder(
        company,
        job_title
    )

    report = {

        "generated_on":
            datetime.now().isoformat(),

        "company":
            company,

        "job_title":
            job_title,

        "overall_score":
            decision.overall_score,

        "confidence":
            decision.confidence,

        "decision":
            decision.decision,

        "priority":
            decision.priority,

        "automation":
            decision.automation_level,

        "resume_strategy":
            decision.resume_strategy,

        "cover_letter_strategy":
            decision.cover_letter_strategy,

        "application_strategy":
            decision.application_strategy,

        "recommendations":
            decision.recommendations,

        "employer": {

            "company":
                employer.company,

            "industry":
                employer.industry,

            "overall_score":
                employer.overall_score,

            "strengths":
                employer.strengths,

            "risks":
                employer.risks

        }

    }

    filename = folder / "career_report.json"

    with open(filename, "w", encoding="utf-8") as f:

        json.dump(
            report,
            f,
            indent=4,
            ensure_ascii=False
        )

    return str(filename)