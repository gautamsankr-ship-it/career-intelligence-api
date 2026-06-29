from docx import Document
from pathlib import Path


OUTPUT_FOLDER = Path("app/output")
OUTPUT_FOLDER.mkdir(parents=True, exist_ok=True)


def generate_resume(data):

    filename = f"{data.candidate_name.replace(' ', '_')}_Resume.docx"

    file_path = OUTPUT_FOLDER / filename

    document = Document()

    document.add_heading(data.candidate_name, level=1)

    document.add_heading("Optimized Resume", level=2)

    document.add_paragraph(data.resume_text)

    document.save(file_path)

    return {
    "path": str(file_path),
    "filename": filename
}