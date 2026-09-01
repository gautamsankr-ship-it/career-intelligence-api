"""Regression for Task 21.10A: generate_resume_docx/generate_cover_letter_docx used to
dump Markdown text verbatim into Word paragraphs (literal #, ##, **), and any
non-list skill/software value was stringified with str(dict). These tests exercise
the fix (_write_markdown_to_docx / _add_inline_runs) directly and end-to-end,
never touching the production applications/ directory."""

import docx
from pypdf import PdfReader

from app.services import docx_service
from app.services.docx_service import (
    _write_markdown_to_docx,
    generate_cover_letter_docx,
    generate_cover_letter_pdf,
    generate_resume_docx,
    generate_resume_pdf,
)


SAMPLE_MARKDOWN = """# Jane Candidate

**Senior Accountant**

jane@example.test | +1 5551234

---

## Professional Summary

**Target Position:** Senior Accountant

**Core Focus Areas**
- Tax planning
- Financial reporting

## Core Competencies

**Software**
- **accounting:** Xero, MYOB
- **analytics:** Excel, Power BI
"""


def _render(markdown_text):
    document = docx.Document()
    _write_markdown_to_docx(document, markdown_text)
    return document


def _paragraph_texts(document):
    return [p.text for p in document.paragraphs]


def test_headings_render_as_styled_paragraphs_not_literal_hashes():
    """Task 21.31: headings are now direct bold/sized paragraphs (not
    Word's built-in Heading N style, which offered no clean way to combine
    keep-with-next + a subtle rule + restrained caps) -- name unchanged
    case, ## section headings rendered in restrained caps as a pure
    display transform (never changing the underlying section label text
    itself, which resume_generator.py already writes in Title Case)."""
    document = _render(SAMPLE_MARKDOWN)
    texts = _paragraph_texts(document)
    assert "Jane Candidate" in texts
    assert "PROFESSIONAL SUMMARY" in texts
    assert "CORE COMPETENCIES" in texts
    for text in texts:
        assert not text.lstrip().startswith("#")

    name_paragraph = next(p for p in document.paragraphs if p.text == "Jane Candidate")
    assert name_paragraph.runs[0].bold is True
    assert name_paragraph.runs[0].font.size.pt == 20

    section_paragraph = next(p for p in document.paragraphs if p.text == "PROFESSIONAL SUMMARY")
    assert section_paragraph.runs[0].bold is True


def test_bold_markers_become_real_bold_runs_not_literal_asterisks():
    document = _render(SAMPLE_MARKDOWN)
    texts = _paragraph_texts(document)
    assert not any("**" in text for text in texts)
    bold_runs = [run for p in document.paragraphs for run in p.runs if run.bold]
    assert any(run.text == "Senior Accountant" for run in bold_runs)
    assert any(run.text == "Target Position:" for run in bold_runs)


def test_bullets_render_as_list_paragraphs_without_leading_dash():
    document = _render(SAMPLE_MARKDOWN)
    bullets = [p for p in document.paragraphs if p.style.name == "List Bullet"]
    bullet_texts = [p.text for p in bullets]
    assert "Tax planning" in bullet_texts
    assert "Financial reporting" in bullet_texts
    assert not any(text.startswith("-") for text in bullet_texts)


def test_horizontal_rule_is_dropped_not_rendered_literally():
    document = _render(SAMPLE_MARKDOWN)
    assert "---" not in _paragraph_texts(document)


def test_structured_dict_bullets_render_readably_not_as_python_repr():
    """The composer-side dict->bullet formatting (already fixed in resume_generator.py)
    produces "- **category:** a, b" lines; confirm the DOCX renderer turns the bold
    category prefix into a real run rather than ever needing str(dict) fallback."""
    document = _render(SAMPLE_MARKDOWN)
    texts = _paragraph_texts(document)
    assert "accounting: Xero, MYOB" in texts
    assert not any(text.startswith("{") or "':" in text for text in texts)


def test_generate_resume_docx_has_no_internal_optimized_resume_label(tmp_path, monkeypatch):
    """Task 21.10C: 'Optimized Resume' is an internal job-search-tool artifact
    and must never appear on an employer-facing document."""
    monkeypatch.setattr(docx_service, "OUTPUT_DIR", tmp_path)
    path = generate_resume_docx(SAMPLE_MARKDOWN, company="Acme Co", job_title="Senior Accountant")
    document = docx.Document(path)
    assert "Optimized Resume" not in _paragraph_texts(document)


def test_generate_resume_docx_end_to_end_is_isolated_and_clean(tmp_path, monkeypatch):
    """Full public entry point, but redirected away from the real applications/
    directory so this test never touches production data (Task 21.8D.3 lesson)."""
    monkeypatch.setattr(docx_service, "OUTPUT_DIR", tmp_path)
    path = generate_resume_docx(SAMPLE_MARKDOWN, company="Acme Co", job_title="Senior Accountant")
    assert str(tmp_path) in path
    document = docx.Document(path)
    texts = _paragraph_texts(document)
    assert not any("**" in t or t.lstrip().startswith("#") for t in texts)
    assert "Jane Candidate" in texts + [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]


def test_whole_line_bold_has_no_literal_asterisks():
    document = _render("**Only Bold Line**")
    assert document.paragraphs[0].text == "Only Bold Line"
    assert document.paragraphs[0].runs[0].bold is True


def test_generate_resume_docx_uses_compact_paragraph_spacing(tmp_path, monkeypatch):
    """Task 21.13 section 5: Word's blank-document default (~10pt space-after,
    ~1.15 line spacing on every paragraph, 1in/1.25in margins) turns a dense,
    mostly single-line-bullet resume into several pages of whitespace rather
    than content. generate_resume_docx must apply tighter, professional-
    resume-appropriate spacing/margins instead of the raw template default."""
    monkeypatch.setattr(docx_service, "OUTPUT_DIR", tmp_path)
    path = generate_resume_docx(SAMPLE_MARKDOWN, company="Acme Co", job_title="Senior Accountant")
    document = docx.Document(path)

    normal_format = document.styles["Normal"].paragraph_format
    assert normal_format.line_spacing == 1.0
    assert normal_format.space_after is not None
    assert normal_format.space_after.pt < 10

    bullet_format = document.styles["List Bullet"].paragraph_format
    assert bullet_format.line_spacing == 1.0

    section = document.sections[0]
    assert section.left_margin.inches <= 1.0
    assert section.right_margin.inches <= 1.0


def test_heading_only_line_has_no_literal_hashes():
    document = _render("### Heading Only")
    assert document.paragraphs[0].text == "Heading Only"
    assert document.paragraphs[0].runs[0].bold is True


def test_section_heading_has_understated_bottom_rule_not_a_drawing_object():
    """The "optional simple horizontal separator" is a paragraph border
    (pure XML, part of the text run stream), never a shape/text box that
    could confuse ATS parsing."""
    document = _render(SAMPLE_MARKDOWN)
    section_paragraph = next(p for p in document.paragraphs if p.text == "PROFESSIONAL SUMMARY")
    pbdr = section_paragraph._p.find(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr"
    )
    assert pbdr is not None


def test_contact_and_meta_lines_are_restrained_not_bold():
    """Contact details and headline/meta lines support the hierarchy but
    must never compete visually with section headings or job titles."""
    document = _render(SAMPLE_MARKDOWN)
    contact_paragraph = next(p for p in document.paragraphs if "jane@example.test" in p.text)
    assert not any(run.bold for run in contact_paragraph.runs)


def test_bullet_style_has_hanging_indent_for_wrapped_lines(tmp_path, monkeypatch):
    monkeypatch.setattr(docx_service, "OUTPUT_DIR", tmp_path)
    path = generate_resume_docx(SAMPLE_MARKDOWN, company="Acme Co", job_title="Senior Accountant")
    bullet_format = docx.Document(path).styles["List Bullet"].paragraph_format
    assert bullet_format.left_indent is not None and bullet_format.left_indent.pt > 0
    assert bullet_format.first_line_indent is not None and bullet_format.first_line_indent.pt < 0


def _pdf_text(path):
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def test_generate_resume_pdf_is_ats_safe_and_matches_docx_content(tmp_path, monkeypatch):
    """Task 21.30 Section 1: the PDF sibling carries real, selectable text
    (extractable by pypdf, unlike an image-only conversion) with the SAME
    factual content as the DOCX -- never regenerated, never rewritten, and
    with no literal Markdown syntax leaking through."""
    monkeypatch.setattr(docx_service, "OUTPUT_DIR", tmp_path)
    docx_path = generate_resume_docx(SAMPLE_MARKDOWN, company="Acme Co", job_title="Senior Accountant")
    pdf_path = generate_resume_pdf(SAMPLE_MARKDOWN, company="Acme Co", job_title="Senior Accountant")

    assert pdf_path.endswith("Resume.pdf")
    pdf_text = _pdf_text(pdf_path)
    assert pdf_text.strip(), "PDF must contain real extractable text, not a rasterized image"
    assert "**" not in pdf_text and "##" not in pdf_text

    docx_texts = _paragraph_texts(docx.Document(docx_path))
    for expected in ("Jane Candidate", "Senior Accountant", "Tax planning", "Financial reporting"):
        assert expected in docx_texts or any(expected in t for t in docx_texts)
        assert expected in pdf_text

    # Task 21.31 section 4: the two formats must share the same heading
    # hierarchy, not just the same body text.
    assert "PROFESSIONAL SUMMARY" in docx_texts
    assert "PROFESSIONAL SUMMARY" in pdf_text


def test_generate_cover_letter_pdf_matches_docx_content(tmp_path, monkeypatch):
    monkeypatch.setattr(docx_service, "OUTPUT_DIR", tmp_path)
    cover_letter = "Dear Hiring Manager,\n\nI am **excited** to apply for this role.\n\n- Strong fit\n- Ready to start\n"
    docx_path = generate_cover_letter_docx(cover_letter, company="Acme Co", job_title="Senior Accountant")
    pdf_path = generate_cover_letter_pdf(cover_letter, company="Acme Co", job_title="Senior Accountant")

    assert pdf_path.endswith("CoverLetter.pdf")
    pdf_text = _pdf_text(pdf_path)
    docx_texts = _paragraph_texts(docx.Document(docx_path))
    assert "excited" in pdf_text and any("excited" in t for t in docx_texts)
    assert "Strong fit" in pdf_text and "Strong fit" in docx_texts
    assert "**" not in pdf_text
