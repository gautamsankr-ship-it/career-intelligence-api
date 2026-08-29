"""Task 21.11: proves ApplicationService.prepare_application() is genuinely
vacancy-driven -- it generates only the materials the vacancy's own text
asks for, using the real production pipeline (ApplicationRequirementService,
ResumeComposer/Generator, CustomResponseGenerator, ApplicationEmailComposer),
with no employer/vacancy name hardcoded anywhere in the logic under test.

Fully hermetic: JobEvaluation is constructed directly (mirroring the existing
pattern in test_application_stages.py/test_preview_evaluation_snapshots.py)
so no OpenAI call happens; all file output is redirected via monkeypatch so
nothing touches the real applications/ or output/resumes/ directories; no
Gmail, no tracker/history, no Answer Vault access at all.
"""

from pathlib import Path
from types import SimpleNamespace

import docx
import pytest

from app.services import docx_service
from app.services.application_email_composer import ApplicationEmailComposer
from app.services.application_requirements import ApplicationRequirementService
from app.services.application_service import ApplicationService, JobEvaluation
from app.services.cover_letter_generator import CoverLetterGenerator
from app.services.custom_response_generator import CustomResponseGenerator
from app.services.resume_composer import ResumeComposer
from app.services.resume_generator import ResumeGenerator
from app.services.resume_optimizer import ResumeOptimizer
from app.services.resume_relevance import extract_vacancy_keywords


def _service() -> ApplicationService:
    """prepare_application() only touches these seven collaborators, none of
    which perform any file I/O on construction. The full ApplicationService()
    constructor additionally builds CareerDecisionEngine/EmployerService/
    ATSEngine/ProfileService/RecruiterReasoningService, which cascade into
    real production-profile/knowledge-base file reads unrelated to this
    method -- bypassing __init__ keeps this test hermetic and fast without
    weakening what's actually under test."""
    service = object.__new__(ApplicationService)
    service.requirement_service = ApplicationRequirementService()
    service.resume_strategy_engine = ResumeOptimizer()
    service.resume_composer = ResumeComposer()
    service.resume_generator = ResumeGenerator()
    service.cover_letter_generator = CoverLetterGenerator()
    service.custom_response_generator = CustomResponseGenerator()
    service.email_composer = ApplicationEmailComposer()
    return service


PROFILE = {
    "candidate": {"full_name": "Jane Candidate", "email": "jane@example.test"},
    "experience": {"years": 15},
    "professional_summary": {"headline": "Chartered Accountant with 15+ years of experience."},
    "employment_history": [
        {
            "company": "Australian Accounting Firm",
            "position": "Offshore Accounting Manager",
            "responsibilities": ["Management Accounting", "Financial Reporting", "Australian Tax"],
            "technologies": ["Xero", "MYOB", "QuickBooks"],
        },
    ],
}


def _evaluation(job_description: str, job_analysis: dict | None = None) -> JobEvaluation:
    return JobEvaluation(
        profile=PROFILE,
        job_analysis=job_analysis or {"company": "EnVision Partners", "job_title": "Tax & Business Advisory Accountant",
                                        "required_skills": ["Xero", "MYOB"], "finance_domains": ["Management Accounting"]},
        employer=SimpleNamespace(),
        career_decision=SimpleNamespace(decision="AUTO_APPLY", overall_score=79.0),
        ats_result={
            "ats_score": {"overall_score": 71.2, "grade": "C"},
            "keyword_summary": {"matched": [], "partial": [], "missing": []},
        },
        screening_decision="AUTO_APPLY",
        job_description=job_description,
    )


@pytest.fixture(autouse=True)
def _isolate_output_dirs(tmp_path, monkeypatch):
    monkeypatch.setattr(docx_service, "OUTPUT_DIR", tmp_path / "applications")
    real_evidence_library = Path.cwd() / "app" / "data" / "candidate_evidence_library.json"
    monkeypatch.chdir(tmp_path)
    (tmp_path / "output" / "resumes").mkdir(parents=True, exist_ok=True)
    # ApplicationService() also constructs ProfileService(), which reads the
    # real candidate profile via a CWD-relative path. prepare_application()
    # itself never calls it (it uses evaluation.profile directly), so a
    # minimal stub is enough to satisfy construction without touching the
    # real production profile file.
    profile_dir = tmp_path / "app" / "data"
    profile_dir.mkdir(parents=True, exist_ok=True)
    (profile_dir / "master_candidate_profile.json").write_text("{}", encoding="utf-8")
    # The real (read-only copied) evidence library, so PROFILE's
    # "Australian Accounting Firm" entry below genuinely exercises the
    # Task 21.11 Addendum evidence-library enrichment path end-to-end.
    (profile_dir / "candidate_evidence_library.json").write_text(
        real_evidence_library.read_text(encoding="utf-8"), encoding="utf-8"
    )


ENVISION_STYLE_TEXT = (
    "Tax & Business Advisory Accountant at EnVision Partners.\n\n"
    "How to Apply\n"
    "To apply, please send an email to Sarah at hr@envision.com.au\n\n"
    "Subject line:\nEnVision - Your Name\n\n"
    "Attach your resume and answer the following question in 200 words or less:\n"
    "Tell us what you can bring to EnVision Partners and why you believe you'd be a great fit for our team.\n"
    "Applications that do not follow these instructions will not be considered."
)


def test_envision_pattern_generates_resume_custom_response_and_email_but_no_cover_letter():
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)

    plan = service.prepare_application(evaluation)

    assert plan.requirements.resume_required is True
    assert plan.requirements.cover_letter_required is False
    assert len(plan.requirements.custom_responses) == 1
    assert plan.requirements.custom_responses[0].max_words == 200
    assert plan.requirements.email_required is True
    assert plan.requirements.recipient_email == "hr@envision.com.au"

    assert plan.resume_docx_path is not None
    assert plan.cover_letter_docx_path is None
    assert len(plan.custom_responses) == 1
    assert len(plan.custom_responses[0].split()) <= 200
    assert plan.email_body is not None
    assert plan.email_recipient == "hr@envision.com.au"


def test_envision_pattern_subject_placeholder_resolved_to_candidate_name():
    """The vacancy's stated subject line ("EnVision - Your Name") is a
    template, not literal text -- the generic "your name" placeholder must
    be substituted with the candidate's actual resolved name, never sent
    verbatim."""
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)

    plan = service.prepare_application(evaluation, candidate_name="Jane Candidate")

    assert plan.email_subject is not None
    assert "your name" not in plan.email_subject.lower()
    assert "Jane Candidate" in plan.email_subject


def test_resume_produced_is_employer_ready_no_ats_diagnostics_no_internal_labels():
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    document = docx.Document(plan.resume_docx_path)
    texts = [p.text for p in document.paragraphs if p.text.strip()]
    for forbidden in ("ATS Optimization Summary", "Optimized Resume", "Target Position", "**"):
        assert not any(forbidden in t for t in texts)


def test_cv_and_cover_letter_pattern_generates_both():
    service = _service()
    evaluation = _evaluation("Please submit your CV and cover letter to jobs@example.com.")
    plan = service.prepare_application(evaluation)

    assert plan.requirements.resume_required is True
    assert plan.requirements.cover_letter_required is True
    assert plan.resume_docx_path is not None
    assert plan.cover_letter_docx_path is not None
    assert not plan.custom_responses


def test_resume_only_pattern_generates_only_resume():
    service = _service()
    evaluation = _evaluation("Upload your resume through our portal.")
    plan = service.prepare_application(evaluation)

    assert plan.resume_docx_path is not None
    assert plan.cover_letter_docx_path is None
    assert not plan.custom_responses
    assert plan.email_body is None


def test_ambiguous_vacancy_requires_human_review_and_generates_nothing():
    service = _service()
    evaluation = _evaluation("Please apply through our careers portal.")
    plan = service.prepare_application(evaluation)

    assert plan.requirements.needs_human_review is True
    assert plan.resume_docx_path is None
    assert plan.cover_letter_docx_path is None
    assert not plan.custom_responses
    assert plan.email_body is None
    assert plan.manifest["human_review_required"] is True


def test_manifest_is_reviewable_and_reflects_generated_artifacts():
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    assert plan.manifest["requirements"]["resume"] is True
    assert plan.manifest["requirements"]["cover_letter"] is False
    assert plan.manifest["requirements"]["custom_response"] is True
    assert plan.manifest["requirements"]["custom_response_max_words"] == 200
    assert plan.manifest["requirements"]["email"] is True
    assert "resume" in plan.manifest["generated_artifacts"]
    assert "cover_letter" not in plan.manifest["generated_artifacts"]
    assert plan.manifest["evidence"]  # evidence is preserved, not discarded


def test_preparation_never_touches_tracker_gmail_or_answer_vault():
    """No import of ApplicationHistoryService/GmailService/ApplicationAnswerVault
    is even reachable from prepare_application -- this test documents that
    boundary explicitly rather than relying on absence of evidence."""
    import app.services.application_service as module
    source = Path(module.__file__).read_text(encoding="utf-8")
    for forbidden in ("GmailService", "ApplicationHistoryService", "ApplicationAnswerVault"):
        assert forbidden not in source


# Task 21.12: the candidate has since confirmed the Trident metrics/Project
# Everest as VERIFIED, so they may now legitimately appear (PROFILE's single
# employment entry matches the evidence library's Trident alias). What must
# still never appear is a claim STRENGTHENED beyond its verified wording, or
# a fact belonging to an employer/venture not present in this PROFILE at all
# (GSN/board/ventures sections are absent here, so their facts should never
# leak in regardless).
FORBIDDEN_STRENGTHENED_CLAIMS = (
    "17 years", "single-handedly", "executed a merger", "led the entire IFRS",
    "Chief Investment Officer", "40 professionals", "$140 million", "Liberty Holdings",
)


def test_evidence_library_enriches_resume_with_verified_facts_only():
    """PROFILE's employment entry (company='Australian Accounting Firm')
    matches the real evidence library's Trident alias, so the resume should
    gain richer VERIFIED facts (SMSF/CAS360, and the now-confirmed 145%/
    Project Everest/12-employee facts) while never strengthening a claim or
    leaking facts that belong to a different, absent employer."""
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    document = docx.Document(plan.resume_docx_path)
    text = "\n".join(p.text for p in document.paragraphs)

    assert "SMSF" in text or "CAS360" in text
    assert "145%" in text
    assert "12 employees" in text
    for marker in FORBIDDEN_STRENGTHENED_CLAIMS:
        assert marker not in text


def test_custom_response_never_surfaces_forbidden_strengthened_claims():
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    response_text = plan.custom_responses[0]
    for marker in FORBIDDEN_STRENGTHENED_CLAIMS:
        assert marker not in response_text


def test_manifest_carries_quality_gate_and_evidence_gaps():
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    assert "quality_gate" in plan.manifest
    assert plan.manifest["quality_gate"]["resume"]["passed"] is True
    assert plan.manifest["quality_gate"]["custom_response"]["passed"] is True
    assert plan.manifest["quality_gate"]["email"]["passed"] is True
    assert "evidence_gaps" in plan.manifest


def test_email_falls_back_to_evidence_grounded_fit_summary_without_custom_response():
    """Email required, no employer-specified custom response to embed --
    the email should still gain a short, evidence-grounded fit summary
    (Task 21.11 Addendum section 13) rather than a bare greeting."""
    service = _service()
    evaluation = _evaluation("Please email your resume to jobs@example.com.")
    plan = service.prepare_application(evaluation)

    assert plan.requirements.email_required is True
    assert not plan.requirements.custom_responses
    assert plan.email_body is not None
    # More than just greeting + closing + signature -- a real fit sentence.
    assert len(plan.email_body.split()) > 15
    for marker in FORBIDDEN_STRENGTHENED_CLAIMS:
        assert marker not in plan.email_body


FULL_EVIDENCE_PROFILE = {
    "candidate": {"full_name": "Shankar Gautam", "email": "gautamsankr@gmail.com"},
    "experience": {"years": 15},
    "professional_summary": {"headline": "Chartered Accountant with 15+ years of experience."},
    "employment_history": [
        {"company": "Australian Accounting Firm", "position": "Offshore Accounting Manager",
         "responsibilities": ["Management Accounting"], "technologies": ["Xero"]},
        {"company": "GSN Associates", "position": "Managing Partner",
         "responsibilities": ["Audit"], "technologies": []},
    ],
    "board_positions": [
        {"organization": "Prabhu Mahalaxmi Life Insurance Limited", "role": "Board of Directors",
         "responsibilities": [], "achievements": []},
    ],
    "entrepreneurship": [
        {"venture": "Sewa360 ERP", "role": "Co-Founder", "achievements": []},
        {"venture": "Liberty Holdings", "role": "Co-Founder", "achievements": []},
    ],
}


def test_employer_facts_never_cross_contaminate_through_the_real_pipeline():
    """Task 21.12 section 11, exercised end-to-end through prepare_application()
    rather than just the evidence-service unit layer: GSN's team-of-40 and
    Trident's team-of-12/145%/Project Everest must land in the right
    employer's own resume section, never the other's."""
    service = _service()
    evaluation = JobEvaluation(
        profile=FULL_EVIDENCE_PROFILE,
        job_analysis={
            "company": "EnVision Partners", "job_title": "Tax & Business Advisory Accountant",
            "required_skills": ["Xero", "BAS", "SMSF"], "finance_domains": ["Australian Tax"],
        },
        employer=SimpleNamespace(),
        career_decision=SimpleNamespace(decision="AUTO_APPLY", overall_score=79.0),
        ats_result={
            "ats_score": {"overall_score": 71.2, "grade": "C"},
            "keyword_summary": {"matched": [], "partial": [], "missing": []},
        },
        screening_decision="AUTO_APPLY",
        job_description=ENVISION_STYLE_TEXT,
    )

    plan = service.prepare_application(evaluation)
    document = docx.Document(plan.resume_docx_path)
    texts = [p.text for p in document.paragraphs if p.text.strip()]

    def section_between(start_markers, end_markers):
        start = next(i for i, t in enumerate(texts) if any(m in t for m in start_markers))
        end = next((i for i in range(start + 1, len(texts)) if any(m in texts[i] for m in end_markers)), len(texts))
        return "\n".join(texts[start:end])

    trident_section = section_between(["Australian Accounting Firm"], ["GSN Associates", "Board Positions"])
    gsn_section = section_between(["GSN Associates"], ["Board Positions", "Entrepreneurship", "Independent Consulting"])

    assert "12 employees" in trident_section or "145%" in trident_section or "Project Everest" in trident_section
    assert "40 professionals" not in trident_section

    assert "40 professionals" in gsn_section
    assert "12 employees" not in gsn_section
    assert "145%" not in gsn_section
    assert "Project Everest" not in gsn_section


# --- Task 21.13: recruiter-facing writing quality, end-to-end -------------

def test_resume_uses_professional_title_capitalization_not_mid_sentence_lowercase():
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    document = docx.Document(plan.resume_docx_path)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Chartered Accountant" in text
    assert "chartered Accountant" not in text
    assert "chartered accountant" not in text


def test_custom_response_uses_professional_title_capitalization():
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    response_text = plan.custom_responses[0]
    assert response_text.startswith("I am a Chartered Accountant")
    assert "chartered Accountant" not in response_text


def test_resume_experience_sentences_are_natural_prose_not_a_field_list():
    """"Led Management Accounting, Financial Reporting, Australian Tax and
    Leadership as Offshore Accounting Manager..." reads like a raw field
    dump. The rewritten sentence should read naturally instead."""
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    document = docx.Document(plan.resume_docx_path)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Led Management Accounting" not in text
    assert "Led Audit, Corporate Finance" not in text
    assert "As Offshore Accounting Manager at Australian Accounting Firm, focused on" in text


def test_resume_no_longer_has_overlapping_skills_sections():
    """Core Focus Areas / Core Skills / Technical Skills / Industry Expertise
    are consolidated into a single Core Competencies section (Task 21.13
    section 4)."""
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    document = docx.Document(plan.resume_docx_path)
    heading_texts = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert heading_texts.count("Core Competencies") == 1
    for forbidden_heading in ("Core Focus Areas", "Technical Skills", "Industry Expertise"):
        assert forbidden_heading not in heading_texts


def test_custom_response_uses_more_of_a_generous_limit_than_a_tight_one():
    """Task 21.12/21.13: with a 200-word allowance and enough relevant
    evidence, the response should use meaningfully more of that allowance
    than a tight limit would allow -- substantive, not padded to a fixed
    band regardless of the fixture's amount of evidence."""
    service = _service()
    evaluation = _evaluation(ENVISION_STYLE_TEXT)
    plan = service.prepare_application(evaluation)

    rich_word_count = len(plan.custom_responses[0].split())
    tight_response = service.custom_response_generator.generate(
        evaluation.profile, max_words=60, employer_name="EnVision Partners",
        vacancy_keywords=extract_vacancy_keywords(evaluation.job_analysis, evaluation.ats_result),
    )
    assert rich_word_count > len(tight_response.split())
    assert rich_word_count <= 200
